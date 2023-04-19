import string
import time
print('Iniciando avalúo ...')
BP = time.time()

from num2words import num2words

# import datetime
from datetime import datetime
import locale
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
# Obtener la fecha actual
fecha_actual = datetime.now()

# Obtener el día, mes y año
dia = fecha_actual.day
mes = fecha_actual.strftime("%B") # %B muestra el nombre del mes en letras
anio = fecha_actual.year

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


document = Document()

# document.add_heading('Document Title', 0)

document.add_paragraph(f"""Bogotá, D. C. {dia} de {mes} de {anio} \n """)

p = document.add_paragraph(f"""Señores\n""")
p.add_run('ELECTROHUILA\n').bold = True
p.add_run('Neiva, Huila\n')
#Variables a capturar
Propietarios = ['JUAN PEDRO']
Cedulas = ['123456']
add_text1 = ''
N_Predio = 'GUAYACANES' ##Nombre del predio
TPredio = 'RURAL' ##Predio en zona rural
Mun = 'Rivera'
TDR = 'PROPIETARIO' ##Calidad del propietario
DAP = 'E.P. N° 1009 del 08 de abril de 2014 de la Notaría Tercera del municipio de Rivera, departamento de Huila, debidamente registrada en la anotación N° 1 del CTL 200-235397.' ##Documetos que acredita propiedad

V_N = 'Los Medios.'      ##Norte
V_S = 'Bajo Pedregal.'##Sur 
V_O = 'Los Medios.'##Oriente
V_OR = 'El Guadual.'## Occidente

POT = 'PBOT'  ##Tipo Ordenamiento Territorial
Ac_POT = 'Acuerdo 004 de 2022'
Nom_POT = 'POR MEDIO DEL CUAL SE APRUEBA EL PLAN BÁSICO DE ORDENAMIENTO TERRITORIAL PBOT DEL MUNICIPIO DE RIVERA (HUILA)'
Art_POT = '23'


i_d = 'd' ##Margen izquierda o derecha

track = 5                                       ## Duraciń trayecto
w_track = num2words(track, lang = 'es')  ## Trayecto en letras

VP = 'Neiva – Rivera' ##Vía Principal al predio
OR_VIA = 'segundo' ##Orden de la via
V_CP = 'Arenoso' #3Via a centro poblado
OV_CP = 'Tercer' ##Orden vía a centro poblado

##1.9.1
lat = 2.79394
lon = -75.25974
alt = 705 

date = '17/07/2023'
# Convertimos la fecha a un objeto datetime
date = datetime.strptime(date, '%d/%m/%Y')
# Formateamos la fecha como un string en el formato deseado
date = date.strftime('%d de %B de %Y')

P = 'Y'
CTL = 'Y'
CUS = 'Y'
EP = 'N'

CC_Num = '416150000000000060079000000000'
CTL_Num = '200-34357'
Date_CTL = '30/05/2021'
Date_CTL = datetime.strptime(Date_CTL, '%d/%m/%Y')
Date_CTL_str = Date_CTL.strftime('%d de %B de %Y')
Date_CUS = '30/04/2021'
Date_CUS = datetime.strptime(Date_CUS, '%d/%m/%Y')
Date_CUS_str = Date_CUS.strftime('%d de %B de %Y')


# lp = len(Propietarios)
if len(Propietarios) > 1:
    add_text = 'de los señores'
    for _ in range(0,len(Propietarios)):
        add_text1 = add_text1 + ',' +Propietarios[_]
else:
    add_text = 'del señor'


# print(Propietarios[1])
# print(add_text1)

p = document.add_paragraph(f"""Atendiendo su amable solicitud, adjunto a la presente le estamos enviando el avalúo comercial practicado al bien inmueble de propiedad {add_text}""")
p.add_run({add_text1}).bold = True
p.add_run(f""";predio ubicado en la vereda Los Vereda""")
p.add_run(f"""VEREDA""").bold = True
p.add_run(f""",jurisdicción del municipio de""")
p.add_run(f"""MUNICIPIO""").blod = True
p.add_run('en el departamento del Huila')

document.add_paragraph(f"""El valor comercial del inmueble ha sido elaborado según los parámetros y criterios de la Resolución No. 620 del 23 de septiembre de 2.008, por medio de la cual se establecieron los procedimientos para la elaboración de los avalúos ordenados dentro del marco de la Ley 388 de 1.997. \n \n Informo expresa y motivadamente que los precios nominales del mercado al que pertenece el inmueble avaluado no han experimentado caídas significativas y duraderas en los últimos cinco (5) años. \n \n Cordialmente""")

p = document.add_paragraph('')
p.add_run(f'CAMILO ANDRÉS LÓPEZ VELASCO \n')
p.add_run(f'Ingeniero Catastral y Geodesta \n').bold = True
p.add_run(f'CAMILO ANDRÉS LÓPEZ VELASCO\n').bold = True
p.add_run(f'M.P 25222-339415 CND\n').bold = True
p.add_run(f'RAA: AVAL – 1015433632 de acuerdo con el autorregulador nacional de Avaluadores ANA.').bold = True

p = document.add_paragraph('')
p.add_run(f'AVALÚO COMERICAL INMUEBLE RURAL\n').bold = True
p.add_run(f'PROPÓSITO DEL AVALUÓ:\n').bold = True
p.add_run(f'Estimar el valor comercial o de mercado del bien inmueble identificado en el Capítulo 1, numeral 1.7 del informe valuatorio, teniendo en cuenta las condiciones económicas reinantes al momento del avalúo y los factores de comercialización que puedan incidir positiva o negativamente en el resultado final.\n \n')
p.add_run(f'DEFINICIÓN DE VALOR COMERCIAL O DE MERCADO:\n').bold = True
p.add_run(f'El valor comercial o de mercado como se utiliza en este informe se define como: la cuantía estimada por la que un bien podría intercambiarse en la fecha de valuación, entre un comprador dispuesto a comprar y un vendedor dispuesto a vender, en una transacción libre tras una comercialización adecuada, en la que las partes hayan actuado con la información suficiente, de manera prudente y sin coacción. \n \n')
p.add_run(f'DERECHO DE PROPIEDAD:\n').bold = True
p.add_run(f'Se considera que el propietario tiene derecho de propiedad completo y absoluto pudiendo disponer y transferir el inmueble con entera libertad.\n \n') 
p.add_run(f'MAYOR Y MEJOR USO:\n').bold = True
p.add_run(f'Se define como: El uso más probable de un bien, físicamente posible, justificado adecuadamente, permitido jurídicamente, financieramente viable y que da como resultado el mayor valor del bien valorado.\n \n')
p.add_run(f'VIGENCIA DEL AVALUÓ:\n').bold = True
p.add_run(f'El presente avalúo tiene vigencia de un año a partir de la fecha de su emisión, siempre y cuando no se presenten circunstancias o cambios inesperados de índole jurídica, técnica, económica o normativa que afecten o modifiquen los criterios analizados. \n \n')

document.add_page_break()


p = document.add_paragraph('')
p.add_run(f'1.1.- CLASE DE AVALÚO:\n').bold = True
p.add_run(f'Comercial.\n')
p.add_run(f'1.2.- TIPO DE INMUEBLE:\n').bold = True	
p.add_run(f'El inmueble objeto de avalúo consiste en un lote de terreno denominado')
p.add_run(f'{N_Predio}').bold = True
p.add_run(f', tipo de predio ')
p.add_run(f'{TPredio}').bold = True
p.add_run(f'segun certificado de tradición y libertad \n')
 
p.add_run(f'1.3. DESTINACIÓN ECONÓMICA ACTUAL DEL INMUEBLE:\n').bold = True
p.add_run(f'Actualmente el inmueble está siendo explotado en el área de franja de servidumbre, sin embargo, el sector en el que se encuentra es de zona agropecuaria, donde se identifican varios lotes con actividad económica de tipo pecuaria y explotación turística.\n')
p.add_run(f'1.4. SOLICITANTE DEL AVALUÓ:\n').bold = True	
p.add_run(f'Electrohuila.\n')
p.add_run(f'1.5. PROPIETARIO DEL INMUEBLE:\n').bold = True	
resultado = ''
add_text3 = []

##Pendiente para varios propietarios
if len(Propietarios) == 1:
    p.add_run(f'{Propietarios[0]} identificado con cédula de ciudadanía N° {Cedulas[0]}')
else: 
    pass

# for _ in range(0, len(Propietarios)):
    
#     # print(Cedulas[0])
#     # print(Cedulas[1])
#     add_text3.append(f'{Propietarios[_]} identificado con cédula de ciudadanía N° {Cedulas[_]} y ')
    # print(f'{Propietarios[_]} identificado con cédula de ciudadanía N° {Cedulas[_]} y')
# add_text31 = ''.join(add_text3)
# p.add_run(f"""{add_text31}""")

p.add_run(f'en calidad de {TDR} de conformidad, con la {DAP}.\n \n')

p.add_run(f'1.6. NOMBRE DEL INMUEBLE:\n').bold = True
p.add_run(f'Predio denominado ')
p.add_run(f'{N_Predio} \n').bold = True

p.add_run(f'1.7. DIRECCIÓN DEL INMUEBLE:\n').bold = True
p.add_run(f'{N_Predio} \n')
p.add_run(f'1.7.1 PARAJE O VEREDA:\n').bold = True

##Falta realizar intersección de la vereda
p.add_run(f'VEREDA (De acuerdo con información geográfica del DANE y Certificado de Uso del Suelo) \n')
p.add_run(f'1.7.2 MUNICIPIO:\n').bold = True
p.add_run(f' {Mun}\n')

p.add_run(f'1.7.3 MUNICIPIO:\n').bold = True
p.add_run(f' Huila \n')

p.add_run(f'1.8 DELIMITACIÓN DEL SECTOR:\n').bold = True
p.add_run(f'Norte: Vereda {V_N}\n')
p.add_run(f'Sur: Vereda {V_S}\n')
p.add_run(f'Occidente: Vereda {V_O}\n')
p.add_run(f'Oriente: Vereda {V_OR}\n')

p.add_run(f'1.9 LOCALIZACIÓN GEOGRÁFICA DEL INMUEBLE: :\n').bold = True

p.add_run(f'El inmueble denominado {N_Predio}, materia de este avalúo, es un predio rural cuyos usos se establecen en el articulo {Art_POT} del {POT} del municipio de {Mun}.')
p.add_run(f'Predio que se encuentra ubicado al costado {i_d} de la vía veredal.')
dist = '10'

track = 5
w_track = num2words(track, lang = 'es')
 
p.add_run(f'El inmueble sujeto a avalúo se encuentra ubicado a {dist} km aproximadamente de la cabecera municipal de Rivera por las vías secundarias y terciarias del municipio; el trayecto puede ser de {w_track} ({track}) minutos \n')
p.add_run(f'El sector de influencia en de zona agroturistica donde se identifican varios lotes con actividad económica de tipo agropecuario, pecuario y explotación turística.\n')
p.add_run(f'La infraestructura vial y de servicios públicos, funcionan satisfactoriamente; la vía principal de la zona de influencia e inmediatamente circundante corresponde a la {VP}, esta se encuentra en pavimento, vía de segundo orden de acuerdo con el {POT} del municipio de {Mun}, en buen estado de conservación, adicionalmente cuenta con la vía al centro poblado de la vereda {V_CP}, vía de tercer nivel según el {POT} de {Mun}. El sector cuenta con redes de acueducto veredal y de riego, pozo séptico y energía Eléctrica. \n')

p.add_run(f'1.9.1. COORDENADAS GEOGRÁFICAS DEL ACCESO PRINCIPAL DEL INMUEBLE:\n').bold = True
lat = 2.79394
lon = -75.25974
alt = 705 
p.add_run(f'latitud:').bold = True
p.add_run(f'{lat}°N \n')
p.add_run(f'Longitud:').bold = True
p.add_run(f'{lon}°O \n')
p.add_run(f'Altura:').bold = True
p.add_run(f'{alt}msnm \n')

p.add_run(f'1.10. VECINDARIO INMEDIATO: \n').bold = True
p.add_run(f'El sector de influencia se halla conformado por pequeños y medianos lotes explotados para actividad económica agroturística y lotes sin explotación en la zona agropecuaria del municipio de {Mun}.')

p.add_run(f'1.11. INFRAESTRUCTURA VIAL: \n').bold = True
p.add_run(f'La vía más importante de la zona es la que conduce al Centro Poblado {V_CP}, la cual es vía de {OV_CP} según {POT} de {Mun} y se encuentra en buen estado de conservación.\n')

p.add_run(f'1.12. INFRAESTRUCTURA DE SERVICIOS PÚBLICOS: \n').bold = True
p.add_run(f'El sector cuenta con las redes instaladas de los servicios públicos de acueducto, alcantarillado y energía eléctrica (Rurales).\n')

p.add_run(f'1.13. TRANSPORTE PÚBLICO: \n').bold = True
p.add_run(f'El sector cuenta con transporte intermunicipal por la vía descrita anteriormente y rutas rurales (camperos y motos).\n')

p.add_run(f'1.14. ESTRATIFICACIÓN SOCIOECONÓMICA DE LA ZONA: \n').bold = True
p.add_run(f'No aplica para el presente caso, teniendo en cuenta que, el estrato aplica a los inmuebles de uso residencial.\n')

p.add_run(f'1.15 FECHA DE LA VISITA: \n').bold = True
p.add_run(f'{date} \n')

p.add_run(f'2. DOCUMENTOS SUMINISTRADOS PARA EL AVALÚO: \n').bold = True
P = 'Y'
CTL = 'Y'
CUS = 'Y'
EP = 'N'

CTL_Num = '200-34357'
Date_CTL = '30/04/2021'
Date_CUS = '30/04/2021'


if P == 'Y':
    # q = document.add_paragraph('Plano de servidumbre aportado por ElectroHuila', style='List Bullet')
    p.add_run(f'        Plano de servidumbre aportado por ElectroHuila \n')
else: 
    pass
if CTL == 'Y':
    # q = document.add_paragraph(f'Certificado de tradición y libertad No {CTL_Num} de fecha {Date_CTL_str}', style='List Bullet')
    p.add_run(f'        Certificado de tradición y libertad No {CTL_Num} de fecha {Date_CTL_str} \n')
else: 
    pass
if CUS == 'Y':
    # q = document.add_paragraph(f'Certificado de uso del suelo de fecha {Date_CUS_str}', style='List Bullet')
    p.add_run(f'        Certificado de uso del suelo de fecha {Date_CUS_str} \n')
else:
    pass
p.add_run(f'        {DAP}\n')


p.add_run(f'3. INFORMACIÓN JURÍDICA \n').bold = True
p.add_run(f'3.1. PROPIETARIO DEL INMUEBLE:	 \n').bold = True

resultado = ''
add_text3 = []

##Pendiente para varios propietarios
if len(Propietarios) == 1:
    p.add_run(f'{Propietarios[0]} identificado con cédula de ciudadanía N° {Cedulas[0]}')
else: 
    pass

# for _ in range(0, len(Propietarios)):
    
#     # print(Cedulas[0])
#     # print(Cedulas[1])
#     add_text3.append(f'{Propietarios[_]} identificado con cédula de ciudadanía N° {Cedulas[_]} y ')
    # print(f'{Propietarios[_]} identificado con cédula de ciudadanía N° {Cedulas[_]} y')
# add_text31 = ''.join(add_text3)
# p.add_run(f"""{add_text31}""")

p.add_run(f'en calidad de {TDR} de conformidad, con la {DAP}.\n \n')
p.add_run(f'3.2. TITULO DE PROPIEDAD Y/O ESCRITURA PÚBLICA: \n').bold = True
p.add_run(f'{DAP} \n')

p.add_run(f'3.3. MATRICULA INMOBILIARIA: \n').bold = True
p.add_run(f'{CTL_Num} \n')

p.add_run(f'3.4. CEDULA CATASTRAL: \n').bold = True
p.add_run(f'{CC_Num} \n')

p.add_run(f'4. NORMATIVIDAD VIGENTE: \n').bold = True
p.add_run(f'La Reglamentación de Usos del Suelo en el Municipio de {Mun}, está contenida en el {Ac_POT} {Nom_POT} que determina el sector como suelo rural y cuyos usos están especificados en el Ártículo {Art_POT} del {POT} del municipio de {Mun}.  \n')

##Anexarr imagen cruce de capas POT

p.add_run(f'Imagen: Normatividad urbanística frente a predio objeto de avalúo. Fuente: Cartografía de {POT} de {Mun}, Huila')

##Anexar imagen Artículo 45 

p.add_run(f'Imagen: LEY 1228 DE 2008 (julio 16) Por la cual se determinan las fajas mínimas de retiro obligatorio o áreas de exclusión, para las carreteras del sistema vial nacional, se crea el Sistema Integral Nacional de Información de Carreteras y se dictan otras disposiciones. \n')

p.add_run('Artículo 18.').bold = True
p.add_run(f' Grávense con la servidumbre legal de conducción de energía eléctrica los predios por los cuales deban pasar las líneas respectivas. \n')
p.add_run(f'Ley 56 de 1981 (Por la cual se dictan normas sobre obras públicas de generación eléctrica, y acueductos, sistemas de regadío y otras y se regulan las expropiaciones y servidumbres de los bienes afectados por tales obras) en sus artículos 25 al 32: \n').bold = True
p.add_run('ARTÍCULO 25:').bold = True
p.add_run(f'La servidumbre pública de conducción de energía eléctrica establecida por el artículo 18 de la Ley 126 de 1938, supone para las entidades públicas que tienen a su cargo la construcción de centrales generadoras, líneas de interconexión, transmisión y prestación del servicio público de distribución de energía eléctrica, la facultad de pasar por los predios afectados, por vía aérea subterránea o superficial, las líneas de transmisión y distribución del fluido eléctrico, ocupar las zonas objeto de la servidumbre, transitar por los mismos, adelantar las obras, ejercer la vigilancia, conservación y mantenimiento y emplear los demás medios necesarios para su ejercicio. \n')
p.add_run(f'ARTÍCULO 26:').bold = True
p.add_run(f'En el trazado de la servidumbre a que se refiere la presente Ley, se atenderá a las exigencias técnicas de la obra. \n')

p.add_run('ARTÍCULO 27:').bold = True 
p.add_run(f'Corresponde a la entidad de derecho público que haya adoptado el respectivo proyecto y ordenado su ejecución, promover en calidad de demandante los procesos que sean necesarios para hacer efectivo el gravamen de servidumbre de conducción de energía eléctrica. Ver Decreto Nacional 2024 de 1982 \n')
p.add_run('Sin perjuicio de las reglas generales contendidas en los libros 1 y 2 del Código de Procedimiento Civil, que le serán aplicables en lo pertinente, el proceso de servidumbre de conducción de energía eléctrica se sujetará a las siguientes reglas: \n')

# num_1 = document.add_paragraph()
# numbered_list_style = document.styles.add_style('Numbered List', 1)
# numbered_list_style.base_style = document.styles['List Number']
# numbered_list_style.paragraph_format.left_indent = Inches(0.5)

# num_1 = document.add_paragraph(style='Numbered List')
p.add_run(f'1. A la demanda se adjuntará el plano general en que figure el curso que habrá de seguir la línea objeto del proyecto con la demarcación específica del área, inventario de los daños que se causen, con el estimativo de su valor realizado por la entidad interesada en forma explicada y discriminada, que se adjuntará al acta elaborada al efecto y certificado de tradición y libertad del predio. \n')

p.add_run(f'Es aplicable a este proceso, en lo pertinente, el artículo 19 de la presente Ley.')

p.add_run(f'2. Con la demanda, la entidad interesada pondrá a disposición del juzgado la suma correspondiente al estimativo de la indemnización. \n')
p.add_run(f'3. Una vez admitida la demanda, se correrá traslado de ella al demandado por el término de tres (3) días. \n')
p.add_run(f'4. Si dos (2) días después de proferido el auto que ordena el traslado de la demanda ésta no hubiere podido ser notificada a los demandados, se procederá a emplazarlos en la forma indicada en el inciso 2 del artículo 452 del Código de Procedimiento Civil. \n')
p.add_run(f'5. Sin perjuicio del deber del juez de abstenerse de proferir sentencia de fondo en los casos previstos por la ley, en este proceso no pueden proponerse excepciones. \n')
document.add_paragraph('')

p.add_run('ARTÍCULO 28: ').bold = True
p.add_run(f"""El juez, dentro de las cuarenta y ocho (48) horas siguientes a la presentación de la demanda, practicará una inspección judicial sobre el predio afectado y autorizará la ejecución de las obras, que de acuerdo con el proyecto sean necesarias para el goce efectivo de la servidumbre. \n""")
p.add_run(f'En la diligencia, el juez identificará el inmueble y hará un examen y reconocimiento de la zona objeto del gravamen. \n')

p.add_run('ARTÍCULO 29: ').bold = True
p.add_run('Cuando el demandado no estuviere conforme con el estimativo de los perjuicios, podrá pedir dentro de los cinco (5) días siguientes a la notificación del auto admisorio de la demanda, que por peritos designados por el juez se practique avalúos de los daños que se causen y tasen la indemnización a que haya lugar por la imposición de la servidumbre. Los peritos se nombrarán conforme a lo indicado en el artículo 21 de esta Ley. \n')

p.add_run('ARTÍCULO 30: ').bold = True
p.add_run(f'Al poseedor o tenedor del predio gravado no le es permitido realizar en éste, acto y obra alguna que pueda perturbar, alterar, disminuir, hacer incómodo o peligroso el ejercicio de la servidumbre de conducción de energía eléctrica, tal como ésta haya quedado establecida, según los planos del proyecto respectivo. Si por razón de nuevas circunstancias fuere necesario introducir variaciones en el modo de ejercer la servidumbre, el poseedor o tenedor del predio gravado está obligado a permitirlas, pero quedará a salvo su derecho de exigir la indemnización por los daños que tales variaciones le causen. \n')

p.add_run('ARTÍCULO 31: ').bold = True
p.add_run('Con base en los estimativos, avalúos, inventarios o pruebas que obren en el proceso, el juez dictará sentencia, señalará el monto de la indemnización y ordenará su pago. \n')
p.add_run('Si en la sentencia se fijare una indemnización mayor que la suma consignada, la entidad demandante deberá consignar la diferencia en favor del poseedor o tenedor del predio, y desde la fecha que recibió la zona objeto de la servidumbre hasta el momento en que deposite el saldo, reconocerá intereses sobre el valor de la diferencia, liquidados según la tasa de interés bancario corriente en el momento de dictar la sentencia. \n')

p.add_run('ARTÍCULO 32: ').bold = True 
p.add_run('Cualquier vacío en las disposiciones aquí establecidas para el proceso de la imposición de la servidumbre de conducción de energía eléctrica, se llenará con las normas de que habla el Título XXII, Libro 2 del Código de Procedimiento Civil. \n')

p.add_run('Ley 142 de 1994 (por la cual se establece el régimen de los servicios públicos domiciliarios y se dictan otras disposiciones.) en sus artículos 56 y 57: \n').bold = True
p.add_run('ARTÍCULO 56: ').bold = True
p.add_run('DECLARATORIA DE UTILIDAD PÚBLICA E INTERÉS SOCIAL PARA LA PRESTACIÓN DE SERVICIOS PÚBLICOS.').bold = True
p.add_run('Declárase de utilidad pública e interés social la ejecución de obras para prestar los servicios públicos y la adquisición de espacios suficientes para garantizar la protección de las instalaciones respectivas. Con ambos propósitos podrán expropiarse bienes inmuebles. \n')
p.add_run('ARTÍCULO 57: FACULTAD DE IMPONER SERVIDUMBRES, HACER OCUPACIONES TEMPORALES Y REMOVER OBSTÁCULOS.').bold = True
p.add_run('Cuando sea necesario para prestar los servicios públicos, las empresas podrán pasar por predios ajenos, por una vía aérea, subterránea o superficial, las líneas, cables o tuberías necesarias; ocupar temporalmente las zonas que requieran en esos predios; remover los cultivos y los obstáculos de toda clase que se encuentren en ellos; transitar, adelantar las obras y ejercer vigilancia en ellos; y, en general, realizar en ellos todas las actividades necesarias para prestar el servicio. El propietario del predio afectado tendrá derecho a indemnización de acuerdo con los términos establecidos en la Ley 56 de 1981, de las incomodidades y perjuicios que ello le ocasione. \n')
p.add_run('Las líneas de transmisión y distribución de energía eléctrica y gas combustible, conducciones de acueducto, alcantarillado y redes telefónicas*, podrán atravesar los ríos, caudales, líneas férreas, puentes, calles, caminos y cruzar acueductos, oleoductos, y otras líneas o conducciones. La empresa interesada, solicitará el permiso a la entidad pública correspondiente; si no hubiere ley expresa que indique quien debe otorgarlo, lo hará el municipio en el que se encuentra el obstáculo que se pretende atravesar.')

p.add_run('CÓDIGO CIVIL (Ley 57 de 1887, art. 4o. Con arreglo al artículo 52 de la Constitución de la República, declárase incorporado en el Código Civil el Título III (arts. 19-52) de la misma Constitución) en sus artículos 879 al 890:')
p.add_run('ARTICULO 879. CONCEPTO DE SERVIDUMBRE.').bold = True
p.add_run('Servidumbre predial o simple servidumbre, es un gravamen impuesto sobre un predio, en utilidad de otro predio de distinto dueño.\n')

p.add_run('ARTICULO 880. SERVIDUMBRES ACTIVAS Y PASIVAS').bold = True
p.add_run('Se llama predio sirviente el que sufre el gravamen, y predio dominante el que reporta la utilidad. \n')
p.add_run('Con respecto al predio dominante, la servidumbre se llama activa, y con respecto al predio sirviente, se llama pasiva. \n')

p.add_run('ARTICULO 881. SERVIDUMBRES CONTINUAS Y DISCONTINUAS.').bold = True
p.add_run('Servidumbre continua es la que se ejerce o se puede ejercer continuamente, sin necesidad de un hecho actual del hombre, como la servidumbre de acueducto por un canal artificial que pertenece al predio dominante; y servidumbre discontinua la que se ejerce a intervalos más o menos largos de tiempo y supone un hecho actual del hombre, como la servidumbre de tránsito. \n')

p.add_run('ARTICULO 882. SERVIDUMBRES POSITIVAS, NEGATIVAS, APARENTES E INAPARENTES.').bold = True
p.add_run('Servidumbre positiva; es, en general, la que sólo impone al dueño del predio sirviente la obligación de dejar hacer, como cualquiera de las dos anteriores; y negativa, la que impone al dueño del predio sirviente la prohibición de hacer algo, que sin la servidumbre le sería lícito, como la de no poder elevar sus paredes sino a cierta altura. \n')
p.add_run('Las servidumbres positivas imponen a veces al dueño del predio sirviente la obligación de hacer algo, como la del artículo 900. \n')
p.add_run('Servidumbre aparente es la que está continuamente a la vista, como la del tránsito, cuando se hace por una senda o por una puerta especialmente destinada a él; e inaparente la que no se conoce por una señal exterior, como la misma de tránsito, cuando carece de estas dos circunstancias y de otras análogas. \n')

p.add_run('ARTICULO 884. PERMANENCIA E INALTERABILIDAD DE LAS SERVIDUMBRES').bold = True
p.add_run('Dividido el predio sirviente no varía la servidumbre que estaba constituida en él, y deben sufrirla aquél o aquéllos a quienes toque la parte en que se ejercía. Así, los nuevos dueños del predio que goza de una servidumbre de tránsito no pueden exigir que se altere la dirección, forma, calidad o anchura de la senda o camino destinado a ella. \n')
p.add_run('ARTICULO 885. DERECHO A LOS MEDIOS PARA USAR LA SERVIDUMBRE.').bold = True 
p.add_run('El que tiene derecho a una servidumbre, lo tiene igualmente a los medios necesarios para ejercerla. Así, el que tiene derecho de sacar agua de una fuente, situada en la heredad vecina, tiene el derecho de tránsito para ir a ella, aunque no se haya establecido expresamente en el título. \n')
p.add_run('ARTICULO 886. DERECHO DE REALIZAR OBRAS INDISPENSABLES PARA USAR LA SERVIDUMBRE.').bold = True 
p.add_run('El que goza de una servidumbre puede hacer las obras indispensables para ejercerla; pero serán a su costa, si no se ha establecido lo contrario; y aun cuando el dueño del predio sirviente se haya obligado a hacerlas o repararlas, le será lícito exonerarse de la obligación, abandonando la parte del predio en que deban hacerse o conservarse las obras.\n')
p.add_run('ARTICULO 887. ALTERACIONES EN LA SERVIDUMBRE.').bold = True 
p.add_run('El dueño del predio sirviente no puede alterar, disminuir, ni hacer más incómoda para el predio dominante la servidumbre con que está gravado el suyo.\n \n')
p.add_run('Con todo, si por el transcurso del tiempo llegare a serle más oneroso el modo primitivo de la servidumbre, podrá proponer que se varíe a su costa; y si las variaciones no perjudican al predio dominante, deberán ser aceptadas. \n')

p.add_run('ARTICULO 888. SERVIDUMBRES NATURALES, LEGALES O VOLUNTARIAS.').bold = True
p.add_run('Las servidumbres, o son naturales, que provienen de la natural situación de los lugares, o legales, que son impuestas por la ley, o voluntarias, que son constituidas por un hecho del hombre. \n \n')
p.add_run('ARTICULO 889. REGULACIÓN DE LA SERVIDUMBRE EN EL CÓDIGO DE POLICÍA.').bold = True
p.add_run('Las disposiciones de este título se entenderán sin perjuicio de lo estatuido sobre servidumbres en el Código de Policía o en otras leyes. \n \n')
p.add_run('ARTICULO 890. DIVISIÓN DEL PREDIO DOMINANTE.')
p.add_run('Dividido el predio dominante, cada uno de los nuevos dueños gozará de la servidumbre, pero sin aumentar el gravamen del predio sirviente. \n \n')

p.add_run('REGLAMENTO TÉCNICO DE INSTALACIONES ELÉCTRICAS (RETIE) \n \n').bold = True
p.add_run('ARTÍCULO 13º. DISTANCIAS DE SEGURIDAD \n \n').bold = True 
p.add_run('Para efectos del presente reglamento y teniendo en cuenta que frente al riesgo eléctrico la técnica más efectiva de prevención, siempre será guardar una distancia respecto a las partes energizadas, puesto que el aire es un excelente aislante, en este apartado se fijan las distancias mínimas que deben guardarse entre líneas o redes eléctricas y elementos físicos existentes a lo largo de su trazado (carreteras, edificaciones, piso del terreno destinado a sembrados, pastos o bosques, etc.), con el objeto de evitar contactos accidentales. Las distancias verticales y horizontales que se presentan en las siguientes tablas, se adoptaron de la norma ANSI C2; todas las tensiones dadas en estas tablas son entre fases, para circuitos con neutro puesto a tierra sólidamente y otros circuitos en los que se tenga un tiempo despeje de falla a tierra acorde con el presente reglamento. \n \n') 
p.add_run('Los constructores y en general quienes presenten proyectos a las curadurías, oficinas de planeación del orden territorial y demás entidades responsables de expedir las licencias o permisos de construcción, deben manifestar por escrito que los proyectos que solicitan dicho trámite cumplen a cabalidad con las distancias mínimas de seguridad establecidas en el RETIE.  \n \n')
p.add_run('Es responsabilidad del diseñador de la instalación eléctrica verificar que en la etapa preconstructiva este requisito se cumpla. No se podrá dar la conformidad con el RETIE a instalaciones que violen estas distancias. El profesional competente responsable de la construcción de la instalación o el inspector que viole esta disposición, sin perjuicio de las acciones penales o civiles, debe ser denunciado e investigado disciplinariamente por el consejo profesional respectivo. \n \n')  
p.add_run('El propietario de una instalación que al modificar la construcción viole las distancias mínimas de seguridad, será objeto de la investigación administrativa correspondiente por parte de las entidades de control y vigilancia por poner en alto riesgo de electrocución no sólo a los moradores de la construcción objeto de la violación, sino a terceras personas y en riesgo de incendio o explosión a las edificaciones contiguas. \n \n')  
p.add_run('A menos que se indique lo contrario, todas las distancias de seguridad deben ser medidas de superficie a superficie. Para la medición de distancias de seguridad, los accesorios metálicos normalmente energizados serán considerados como parte de los conductores de línea y las bases metálicas de los terminales del cable o los dispositivos similares, deben ser tomados como parte de la estructura de soporte. La precisión en los elementos de medida no podrá tener un error de más o menos 0,5%. \n \n') 

p.add_run('Para mayor claridad se deben tener en cuenta las notas explicativas, las figuras y las tablas aquí establecidas. \n \n')  

p.add_run('Nota 1:').bold = True
p.add_run('Las distancias de seguridad establecidas en las siguientes tablas, aplican a conductores desnudos. \n')
p.add_run('Nota 2:').bold = True 
p.add_run('En el caso de tensiones mayores a 57,5 kV entre fases, las distancias de aislamiento eléctrico especificadas en las tablas se incrementarán en un 3% por cada 300 m que sobrepasen los 1000 metros sobre el nivel del mar. \n')
p.add_run('Nota 3:').bold = True 
p.add_run('Las distancias verticales se toman siempre desde el punto energizado más cercano al lugar de posible contacto. \n') 
p.add_run('Nota 4:').bold = True 
p.add_run('La distancia horizontal “b” se toma desde la parte energizada más cercana al sitio de posible contacto, es decir, trazando un círculo desde la parte energizada, teniendo en cuenta la posibilidad real de expansión vertical que tenga la edificación y que en ningún momento la red quede encima de la construcción. \n')  
p.add_run('Nota 5:').bold = True 
p.add_run('Si se tiene una instalación con una tensión diferente a las contempladas en el presente reglamento, debe cumplirse el requisito exigido para la tensión inmediatamente superior. \n') 
p.add_run('Nota 6:').bold = True 
p.add_run('Cuando los edificios, chimeneas, antenas o tanques u otras instalaciones elevadas no requieran algún tipo de mantenimiento, como pintura, limpieza, cambio de partes o trabajo de personas cerca de los conductores; la distancia horizontal “b”, se podrá reducir en 0, 6 m.  \n')
p.add_run('Nota 7:').bold = True 
p.add_run('Un techo, balcón o área es considerado fácilmente accesible para los peatones si éste puede ser alcanzado de manera casual a través de una puerta, rampa, ventana, escalera o una escalera a mano permanentemente utilizada por una persona, a pie, alguien que no despliega ningún esfuerzo físico extraordinario ni emplea ningún instrumento o dispositivo especial para tener acceso a éstos. No se considera un medio de acceso a una escalera permanentemente utilizada si es que su peldaño más bajo mide 2,45 m o más desde el nivel del piso u otra superficie accesible fija.  \n')
p.add_run('Nota 8:').bold = True
p.add_run('Si se tiene un tendido aéreo con cable aislado y con pantalla no se aplican estas distancias; tampoco se aplica para conductores aislados para baja tensión. \n')
p.add_run('Nota 9:').bold = True 
p.add_run('En techos metálicos cercanos o en casos de redes de conducción que van paralelas o que cruzan las líneas de media, alta y extra alta tensión, se debe verificar que las tensiones inducidas no generen peligro o no afecten el funcionamiento de otras redes.  \n')
p.add_run('Nota 10:').bold = True
p.add_run('Donde el espacio disponible no permita cumplir las distancias horizontales de la Tabla 13.1 para redes de media tensión, tales como edificaciones con fachadas o terrazas cercanas, la separación se puede reducir hasta en un 30%, siempre y cuando, los conductores, empalmes y herrajes tengan una cubierta que proporcione suficiente rigidez dieléctrica para limitar la probabilidad de falla a tierra, tal como la de los cables cubiertos con tres capas para red compacta. Adicionalmente, deben tener espaciadores y una señalización que indique que es cable no aislado. En zonas arborizadas urbanas se recomienda usar esta tecnología para disminuir las podas.  \n')
p.add_run('Nota11:').bold = True
p.add_run('En general los conductores de la línea de mayor tensión deben estar a mayor altura que los de la de menor tensión. \n \n')


p.add_run('13.1 DISTANCIAS MÍNIMAS DE SEGURIDAD EN ZONAS CON CONSTRUCCIONES \n ').bold = True 
p.add_run('Las distancias mínimas de seguridad que deben guardar las partes energizadas respecto de las construcciones son las establecidas en la Tabla 13.1 del presente reglamento y para su interpretación se debe tener en cuenta la Figura 13.1. \n \n') 

###Incluir dos imágenes 

p.add_run('13.2 DISTANCIAS MÍNIMAS DE SEGURIDAD PARA DIFERENTES LUGARES Y SITUACIONES \n \n')  
p.add_run('En líneas de trasmisión o redes de distribución, la altura de los conductores respecto del piso o de la vía, como lo señalan las Figuras 13.2 y 13.3, no podrá ser menor a las establecidas en la Tabla 13.2. \n')

###Incluir una imágen

p.add_run('14.RESOLUCIÓN NÚMERO 620 DE 2008 (23 Septiembre 2008) \n')
p.add_run('Por la cual se establecen los procedimientos para los avalúos ordenados dentro del marco de la Ley 388 de 1997. \n')
p.add_run('15.Resolución 1092 del 20 septiembre del 2022 de Servidumbres - IGAC: "Por la cual se fijas normas, métodos, parámetros, criterios y procedimientos para la elaboración de avalúos de servidumbres legales y afectaciones transitorias en desarrollo de actividades, obras o proyectos declarados por el legislador como de utilidad pública e interés social." \n').bold = True 
p.add_run('De la Resolución mencionada previamente se tienen en cuenta los siguientes aspectos: \n \n')


p.add_run('Capítulo II \n').bold = True
p.add_run('Cálculo de la indemnización por objeto de la servidumbre \n').bold = True

p.add_run('Artículo 8°.- Factores para determinar la indemnización: ').bold = True

p.add_run('Se establecerán según el grado de afectación que se genere sobre la franja de terreno requerida para la obra de utilidad pública e interés social, los cuales se determinan a continuación: \n')

p.add_run('     1. Factor según tipo de infraestructura \n')
p.add_run('     2. Factor según la clase de suelo \n')

p.add_run('Artículo 9°.- Factor según tipo de infraestructura.').bold = True 
p.add_run('Corresponde a la afectación según la actividad a la que está destinada la obra de utilidad pública e interés social; para lo cual se aplicará a la siguiente clasificación: \n ')

## Incluir tabla

p.add_run('Parágrafo 1: ').bold = True
p.add_run('Cuando el grado de afectación sea total la indemnización corresponderá al 100% del valor comercial del terreno donde se ubica la franja de servidumbre, de conformidad con el artículo 11 de la presente resolución. \n')


p.add_run('Parágrafo 2: ').bold = True
p.add_run('Cuando el grado de afectación es parcial se continuará según lo establecido para la determinación de la valoración del derecho de servidumbre de acuerdo con la infraestructura en los términos establecidos en los artículos 13 y 14. \n')

p.add_run('Parágrafo 3:').bold = True
p.add_run('Para las diferentes industrias se tendrá en cuenta su grado de afectación según su infraestructura instalada o por instalar, la cual puede encontrarse o no en el listado anterior, donde para la infraestructura enterrada o aérea será parcial y para la infraestructura superficial o a nivel de terreno será el 100%. \n')

p.add_run('Parágrafo 4:').bold = True
p.add_run('Independientemente de la obra de utilidad pública e interés social, cuando la infraestructura que se instale sea aérea, la determinación de la valoración del derecho de servidumbre se realizará según los términos establecidos en el artículo 14. \n')


























# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')


# document.add_picture('monty-truth.png', width=Inches(1.25))

# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

# numeracion1 = document.add_paragraph()
# numeracion1.add_run('1. Este es el primer elemento numerado del primer párrafo.')
# numeracion1.add_run('\n2. Este es el segundo elemento numerado del primer párrafo.')

# Agregar un salto de línea
# document.add_paragraph('')

# Agregar el segundo párrafo con numeración
# numbered_list_style = document.styles.add_style('Numbered List', 1)
# numbered_list_style.base_style = document.styles['List Number']
# numbered_list_style.paragraph_format.left_indent = Inches(0.5)

# # Agregar el primer párrafo con numeración
# numeracion1 = document.add_paragraph(style='Numbered List')
# numeracion1.add_run('Este es el primer elemento numerado del primer párrafo.')
# numeracion1.add_run('Este es el segundo elemento numerado del primer párrafo.')






document.save('demo.docx')




# import docx
# doc.add_paragraph(f'Bogotá, D. C. {dia} de {mes} de {anio}')
# doc.add_paragraph(f"""Señores:
# ELECTROHUILA
# NeivaHuila""")
# doc.save('/home/camilocorredor/DS_P/ETL/Appraisal/Document10.docx')

# import time
# print('Iniciando ITJP ...')
# BP = time.time()
# from ctypes import c_char_p
# from osgeo import ogr, osr

# import shapely.wkb
# from shapely.geometry import Polygon

# import openpyxl
# import os
# os.environ['USE_PYGEOS'] = '0'
# import geopandas as gpd

# import psycopg2

# from shapely.wkb import loads
# from shapely import wkb
# import os

# import binascii
# import matplotlib.pyplot as plt

